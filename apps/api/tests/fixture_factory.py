"""Create harmless binary fixtures for document ingestion tests."""

from __future__ import annotations

import io
from pathlib import Path

import pymupdf
from docx import Document as DocxDocument

FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures"


def write_sample_pdf(path: Path | None = None) -> bytes:
    doc = pymupdf.open()
    page_one = doc.new_page()
    page_one.insert_text((72, 72), "Page one: stability data and shelf life.")
    page_two = doc.new_page()
    page_two.insert_text((72, 72), "Page two: impurity profile summary.")
    data = doc.tobytes()
    doc.close()
    if path is not None:
        path.write_bytes(data)
    return data


def write_sample_docx(path: Path | None = None) -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Paragraph one: nomenclature and chemical name.")
    doc.add_paragraph("Paragraph two: manufacturing process controls.")
    buffer = io.BytesIO()
    doc.save(buffer)
    data = buffer.getvalue()
    if path is not None:
        path.write_bytes(data)
    return data
