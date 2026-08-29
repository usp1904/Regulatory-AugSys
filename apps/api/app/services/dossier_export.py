"""Evidence-based dossier export to PDF, DOCX, and TXT."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import UTC, datetime

import fitz
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from sqlalchemy import func, select
from sqlalchemy.orm import Session, joinedload

from app.core.config import get_settings
from app.models.ctd_section import CtdSection
from app.models.document import Document
from app.models.dossier_export import EXPORT_FORMATS, DossierExport
from app.models.evidence_item import GAP_EXPORT_TYPES, EvidenceItem
from app.schemas.dossier_export import DossierExportResponse, ExportManifest
from app.services.audit import record_audit_event
from app.services.document_storage import persist_original_bytes

WATERMARK = "TRAINING / INTERNAL REVIEW ONLY — NOT A REGULATORY SUBMISSION"
SECTION_INTRO = "The following approved evidence items are recorded for this CTD section."
GAP_BLOCK_HEADER = "Controlled Gap / Confidential Source Required"
MANIFEST_HEADER = "Export Manifest"
DOSSIER_TITLE = "Evidence-Based Dossier Export (Draft Internal Record)"

CONTENT_TYPES = {
    "txt": "text/plain",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


class DossierExportError(ValueError):
    """Business rule violation for dossier export."""


@dataclass(frozen=True)
class ExportEvidenceRow:
    evidence_id: int
    evidence_key: str
    evidence_version: int
    ctd_section_code: str | None
    ctd_section_title: str | None
    evidence_type: str
    exact_source_excerpt: str
    normalized_summary: str | None
    source_filename: str | None
    source_document_version: int | None
    page_number: int | None
    source_file_hash: str | None
    is_gap_block: bool


@dataclass(frozen=True)
class ExportSection:
    code: str | None
    title: str | None
    items: list[ExportEvidenceRow]


def ctd_code_sort_key(code: str | None) -> tuple:
    if not code:
        return (9999, "")
    tokens: list[tuple[int, int | str]] = []
    for part in code.split("."):
        for match in re.finditer(r"\d+|[A-Za-z]+", part):
            token = match.group()
            if token.isdigit():
                tokens.append((0, int(token)))
            else:
                tokens.append((1, token))
    return tuple(tokens)  # type: ignore[return-value]


def _load_ctd_titles(db: Session) -> dict[str, str]:
    rows = db.scalars(select(CtdSection)).all()
    return {row.code: row.title for row in rows}


def _next_dossier_version(db: Session, dossier_id: str) -> int:
    current = db.scalar(
        select(func.coalesce(func.max(DossierExport.dossier_version), 0)).where(
            DossierExport.dossier_id == dossier_id
        )
    )
    return int(current or 0) + 1


def collect_approved_evidence(db: Session, dossier_id: str) -> list[ExportEvidenceRow]:
    items = db.scalars(
        select(EvidenceItem)
        .options(joinedload(EvidenceItem.source_document))
        .where(EvidenceItem.dossier_id == dossier_id)
        .where(EvidenceItem.review_status == "APPROVED")
    ).all()
    ctd_titles = _load_ctd_titles(db)
    rows: list[ExportEvidenceRow] = []
    for item in items:
        doc: Document | None = item.source_document
        rows.append(
            ExportEvidenceRow(
                evidence_id=item.id,
                evidence_key=item.evidence_key,
                evidence_version=item.evidence_version,
                ctd_section_code=item.ctd_section_code,
                ctd_section_title=ctd_titles.get(item.ctd_section_code or ""),
                evidence_type=item.evidence_type,
                exact_source_excerpt=item.exact_source_excerpt,
                normalized_summary=item.normalized_summary,
                source_filename=doc.filename if doc else None,
                source_document_version=item.source_document_version,
                page_number=item.page_number,
                source_file_hash=doc.file_hash if doc else None,
                is_gap_block=item.evidence_type in GAP_EXPORT_TYPES,
            )
        )
    rows.sort(key=lambda row: (ctd_code_sort_key(row.ctd_section_code), row.evidence_id))
    return rows


def group_by_section(rows: list[ExportEvidenceRow]) -> list[ExportSection]:
    sections: list[ExportSection] = []
    current_code: str | None = None
    current_items: list[ExportEvidenceRow] = []
    current_title: str | None = None

    for row in rows:
        if row.ctd_section_code != current_code:
            if current_items:
                sections.append(
                    ExportSection(code=current_code, title=current_title, items=current_items)
                )
            current_code = row.ctd_section_code
            current_title = row.ctd_section_title
            current_items = [row]
        else:
            current_items.append(row)

    if current_items:
        sections.append(ExportSection(code=current_code, title=current_title, items=current_items))
    return sections


def _section_heading(section: ExportSection) -> str:
    if section.code and section.title:
        return f"{section.code} — {section.title}"
    if section.code:
        return section.code
    return "Unassigned CTD Section"


def _source_reference(row: ExportEvidenceRow) -> str | None:
    if row.is_gap_block or not row.source_filename:
        return None
    if row.page_number is not None:
        return (
            f"Source: {row.source_filename} "
            f"(version {row.source_document_version}, page {row.page_number})"
        )
    return f"Source: {row.source_filename} (version {row.source_document_version})"


def _statement_text(row: ExportEvidenceRow) -> str:
    if row.normalized_summary:
        return row.normalized_summary
    return row.exact_source_excerpt


def build_manifest(
    *,
    export_id: str,
    dossier_id: str,
    dossier_version: int,
    export_format: str,
    rows: list[ExportEvidenceRow],
) -> ExportManifest:
    document_hashes: dict[str, str] = {}
    for row in rows:
        if row.source_file_hash:
            document_hashes[str(row.evidence_id)] = row.source_file_hash

    unique_hashes: dict[str, str] = {}
    for row in rows:
        if row.source_file_hash and row.source_filename:
            unique_hashes[row.source_filename] = row.source_file_hash

    return ExportManifest(
        export_id=export_id,
        timestamp=datetime.now(UTC),
        dossier_id=dossier_id,
        dossier_version=dossier_version,
        generator_version=get_settings().app_version,
        evidence_ids=[row.evidence_id for row in rows],
        evidence_keys=[row.evidence_key for row in rows],
        document_hashes=unique_hashes,
        export_format=export_format,
        item_count=len(rows),
    )


def _manifest_lines(manifest: ExportManifest) -> list[str]:
    return [
        "",
        "=" * 72,
        MANIFEST_HEADER,
        "=" * 72,
        f"Export ID: {manifest.export_id}",
        f"Timestamp: {manifest.timestamp.isoformat()}",
        f"Dossier ID: {manifest.dossier_id}",
        f"Dossier version: {manifest.dossier_version}",
        f"Generator version: {manifest.generator_version}",
        f"Export format: {manifest.export_format}",
        f"Evidence IDs: {', '.join(str(i) for i in manifest.evidence_ids)}",
        f"Evidence keys: {', '.join(manifest.evidence_keys)}",
        "Document hashes:",
        *[
            f"  - {name}: {digest}"
            for name, digest in sorted(manifest.document_hashes.items())
        ],
        f"Item count: {manifest.item_count}",
    ]


def render_txt(
    dossier_id: str,
    sections: list[ExportSection],
    manifest: ExportManifest,
) -> bytes:
    lines = [
        WATERMARK,
        "",
        DOSSIER_TITLE,
        f"Dossier ID: {dossier_id}",
        "",
    ]
    for section in sections:
        lines.extend(["", "-" * 72, _section_heading(section), "-" * 72, SECTION_INTRO, ""])
        for row in section.items:
            if row.is_gap_block:
                lines.extend(
                    [
                        f">>> {GAP_BLOCK_HEADER} <<<",
                        f"Type: {row.evidence_type}",
                        _statement_text(row),
                        "",
                    ]
                )
            else:
                lines.append(_statement_text(row))
                ref = _source_reference(row)
                if ref:
                    lines.append(ref)
                lines.append("")
    lines.extend(_manifest_lines(manifest))
    return "\n".join(lines).encode("utf-8")


def render_docx(
    dossier_id: str,
    sections: list[ExportSection],
    manifest: ExportManifest,
) -> bytes:
    doc = DocxDocument()
    header = doc.sections[0].header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = WATERMARK
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_heading(DOSSIER_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    watermark_para = doc.add_paragraph(WATERMARK)
    watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    watermark_para.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_paragraph(f"Dossier ID: {dossier_id}")

    for section in sections:
        doc.add_heading(_section_heading(section), level=1)
        doc.add_paragraph(SECTION_INTRO)
        for row in section.items:
            if row.is_gap_block:
                gap_heading = doc.add_paragraph()
                gap_run = gap_heading.add_run(GAP_BLOCK_HEADER)
                gap_run.bold = True
                gap_run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
                type_para = doc.add_paragraph(f"Type: {row.evidence_type}")
                type_para.runs[0].italic = True
                doc.add_paragraph(_statement_text(row))
            else:
                doc.add_paragraph(_statement_text(row))
                ref = _source_reference(row)
                if ref:
                    ref_para = doc.add_paragraph(ref)
                    ref_para.runs[0].font.size = Pt(9)
                    ref_para.runs[0].italic = True

    doc.add_page_break()
    doc.add_heading(MANIFEST_HEADER, level=1)
    for line in _manifest_lines(manifest)[3:]:
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _pdf_add_watermark(page: fitz.Page) -> None:
    rect = page.rect
    page.insert_textbox(
        fitz.Rect(0, rect.height - 40, rect.width, rect.height - 10),
        WATERMARK,
        fontsize=8,
        color=(0.6, 0.6, 0.6),
        align=fitz.TEXT_ALIGN_CENTER,
    )
    page.insert_textbox(
        fitz.Rect(36, 36, rect.width - 36, 72),
        WATERMARK,
        fontsize=10,
        color=(0.85, 0.85, 0.85),
        align=fitz.TEXT_ALIGN_CENTER,
    )


def render_pdf(
    dossier_id: str,
    sections: list[ExportSection],
    manifest: ExportManifest,
) -> bytes:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y = 72
    line_height = 14
    margin = 54
    max_width = page.rect.width - 2 * margin

    def ensure_space(needed: float = line_height) -> None:
        nonlocal page, y
        if y + needed > page.rect.height - 60:
            _pdf_add_watermark(page)
            page = doc.new_page(width=595, height=842)
            y = 72

    def write_line(text: str, *, fontsize: float = 11, gap: float = 4) -> None:
        nonlocal y
        if not text:
            y += gap
            return
        char_width = max(40, int(max_width / (fontsize * 0.5)))
        words = text.split()
        current = ""
        for word in words:
            trial = f"{current} {word}".strip()
            if len(trial) > char_width and current:
                ensure_space(line_height + gap)
                page.insert_text((margin, y), current, fontsize=fontsize)
                y += line_height
                current = word
            else:
                current = trial
        if current:
            ensure_space(line_height + gap)
            page.insert_text((margin, y), current, fontsize=fontsize)
            y += line_height + gap

    write_line(WATERMARK, fontsize=9, gap=8)
    write_line(DOSSIER_TITLE, fontsize=16, gap=8)
    write_line(f"Dossier ID: {dossier_id}", gap=12)

    for section in sections:
        write_line(_section_heading(section), fontsize=13, gap=4)
        write_line(SECTION_INTRO, gap=8)
        for row in section.items:
            if row.is_gap_block:
                write_line(f">>> {GAP_BLOCK_HEADER} <<<", fontsize=11, gap=2)
                write_line(f"Type: {row.evidence_type}", gap=2)
                write_line(_statement_text(row), gap=8)
            else:
                write_line(_statement_text(row), gap=2)
                ref = _source_reference(row)
                if ref:
                    write_line(ref, fontsize=9, gap=8)
                else:
                    y += 4

    write_line("", gap=8)
    for line in _manifest_lines(manifest):
        write_line(line, fontsize=9, gap=2)

    for pdf_page in doc:
        _pdf_add_watermark(pdf_page)

    return doc.tobytes()


def render_dossier_bytes(
    export_format: str,
    dossier_id: str,
    sections: list[ExportSection],
    manifest: ExportManifest,
) -> bytes:
    if export_format == "txt":
        return render_txt(dossier_id, sections, manifest)
    if export_format == "docx":
        return render_docx(dossier_id, sections, manifest)
    if export_format == "pdf":
        return render_pdf(dossier_id, sections, manifest)
    raise DossierExportError(f"Unsupported export format: {export_format}")


def create_dossier_export(
    db: Session,
    dossier_id: str,
    export_format: str,
    actor: str,
) -> DossierExport:
    export_format = export_format.lower()
    if export_format not in EXPORT_FORMATS:
        raise DossierExportError(f"format must be one of {sorted(EXPORT_FORMATS)}")

    rows = collect_approved_evidence(db, dossier_id)
    if not rows:
        raise DossierExportError("No approved evidence items found for dossier")

    sections = group_by_section(rows)
    export_id = DossierExport.new_export_id()
    dossier_version = _next_dossier_version(db, dossier_id)
    manifest = build_manifest(
        export_id=export_id,
        dossier_id=dossier_id,
        dossier_version=dossier_version,
        export_format=export_format,
        rows=rows,
    )
    content = render_dossier_bytes(export_format, dossier_id, sections, manifest)
    filename = f"{dossier_id}_v{dossier_version}.{export_format}"
    stored = persist_original_bytes(filename, content, CONTENT_TYPES[export_format])

    record = DossierExport(
        export_id=export_id,
        dossier_id=dossier_id,
        dossier_version=dossier_version,
        export_format=export_format,
        file_hash=stored.file_hash,
        storage_path=stored.storage_path,
        byte_size=stored.byte_size,
        content_type=stored.content_type,
        manifest_json=manifest.model_dump_json(),
        created_by=actor,
    )
    db.add(record)
    db.flush()
    record_audit_event(
        db,
        event_type="dossier_export",
        actor=actor,
        detail={
            "export_id": export_id,
            "dossier_id": dossier_id,
            "dossier_version": dossier_version,
            "format": export_format,
            "file_hash": stored.file_hash,
            "evidence_ids": manifest.evidence_ids,
            "item_count": manifest.item_count,
        },
    )
    db.commit()
    db.refresh(record)
    return record


def dossier_export_to_response(record: DossierExport) -> DossierExportResponse:
    manifest = ExportManifest.model_validate_json(record.manifest_json)
    return DossierExportResponse(
        id=record.id,
        export_id=record.export_id,
        dossier_id=record.dossier_id,
        dossier_version=record.dossier_version,
        export_format=record.export_format,
        file_hash=record.file_hash,
        byte_size=record.byte_size,
        content_type=record.content_type,
        manifest=manifest,
        created_by=record.created_by,
        created_at=record.created_at,
        download_url=f"/api/v1/dossier-exports/{record.export_id}/download",
    )


def verify_content_excludes_pending_rejected(
    content: bytes,
    *,
    approved_text: str,
    pending_text: str,
    rejected_text: str,
) -> bool:
    text = content.decode("utf-8", errors="ignore").lower()
    approved = approved_text.lower() in text
    pending_absent = pending_text.lower() not in text
    rejected_absent = rejected_text.lower() not in text
    return approved and pending_absent and rejected_absent
