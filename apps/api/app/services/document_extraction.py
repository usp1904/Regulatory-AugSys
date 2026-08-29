"""Extract text from controlled document uploads without modifying originals."""

from __future__ import annotations

import io
from dataclasses import dataclass

import pymupdf
from docx import Document as DocxDocument


@dataclass(frozen=True)
class ExtractionResult:
    status: str
    pages: list[str]
    paragraphs: list[str]
    error: str | None = None

    @property
    def full_text(self) -> str:
        if self.pages:
            return "\n\n".join(self.pages)
        if self.paragraphs:
            return "\n\n".join(self.paragraphs)
        return ""


def extract_pdf_pages(data: bytes) -> ExtractionResult:
    try:
        with pymupdf.open(stream=data, filetype="pdf") as pdf:
            pages = []
            for page in pdf:
                pages.append(page.get_text("text").strip())
        if not pages:
            return ExtractionResult(
                status="EXTRACTION_FAILED",
                pages=[],
                paragraphs=[],
                error="PDF contained no pages",
            )
        return ExtractionResult(status="EXTRACTED", pages=pages, paragraphs=[])
    except Exception as exc:  # noqa: BLE001 — surface extraction failure to caller
        return ExtractionResult(
            status="EXTRACTION_FAILED",
            pages=[],
            paragraphs=[],
            error=str(exc),
        )


def extract_docx_paragraphs(data: bytes) -> ExtractionResult:
    try:
        doc = DocxDocument(io.BytesIO(data))
        paragraphs = [
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        ]
        if not paragraphs:
            return ExtractionResult(
                status="EXTRACTION_FAILED",
                pages=[],
                paragraphs=[],
                error="DOCX contained no non-empty paragraphs",
            )
        return ExtractionResult(status="EXTRACTED", pages=[], paragraphs=paragraphs)
    except Exception as exc:  # noqa: BLE001
        return ExtractionResult(
            status="EXTRACTION_FAILED",
            pages=[],
            paragraphs=[],
            error=str(exc),
        )


def extract_txt_pages(data: bytes) -> ExtractionResult:
    try:
        text = data.decode("utf-8", errors="replace").strip()
        if not text:
            return ExtractionResult(
                status="EXTRACTION_FAILED",
                pages=[],
                paragraphs=[],
                error="TXT file was empty",
            )
        return ExtractionResult(status="EXTRACTED", pages=[text], paragraphs=[])
    except Exception as exc:  # noqa: BLE001
        return ExtractionResult(
            status="EXTRACTION_FAILED",
            pages=[],
            paragraphs=[],
            error=str(exc),
        )


def extract_document(content_type: str, data: bytes) -> ExtractionResult:
    media_type = content_type.split(";")[0].strip().lower()
    if media_type == "application/pdf":
        return extract_pdf_pages(data)
    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_docx_paragraphs(data)
    if media_type == "text/plain":
        return extract_txt_pages(data)
    return ExtractionResult(
        status="UNSUPPORTED_MIME",
        pages=[],
        paragraphs=[],
        error=f"Unsupported media type: {media_type}",
    )
